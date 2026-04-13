from __future__ import annotations

import html
import re
import subprocess
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from .constants import TEXT_EXTENSIONS


class UnsupportedFormatError(Exception):
    pass


class DocumentLoader:
    @staticmethod
    def can_open(path: Path) -> bool:
        return path.suffix.lower() in TEXT_EXTENSIONS | {
            ".docx",
            ".odt",
            ".rtf",
            ".pages",
            ".html",
            ".htm",
        }

    @staticmethod
    def load(path: Path) -> str:
        suffix = path.suffix.lower()

        if suffix in TEXT_EXTENSIONS:
            return DocumentLoader._read_plain_text(path)
        if suffix == ".docx":
            return DocumentLoader._read_docx(path)
        if suffix == ".odt":
            return DocumentLoader._read_odt(path)
        if suffix == ".rtf":
            return DocumentLoader._read_rtf(path)
        if suffix == ".pages":
            return DocumentLoader._read_pages(path)
        if suffix in {".html", ".htm"}:
            return DocumentLoader._read_html(path)

        if suffix == ".pdf":
            return DocumentLoader._read_pdf(path)

        raise UnsupportedFormatError(f"Unsupported format: {path.suffix}")

    @staticmethod
    def _read_plain_text(path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(errors="replace")

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise UnsupportedFormatError(
                "DOCX support requires python-docx. Install it with: pip install python-docx"
            ) from exc

        document = Document(str(path))
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        return "\n\n".join(blocks).strip()

    @staticmethod
    def _read_odt(path: Path) -> str:
        try:
            with zipfile.ZipFile(path) as archive:
                xml_bytes = archive.read("content.xml")
        except Exception as exc:
            raise UnsupportedFormatError("Could not read ODT file") from exc

        root = ET.fromstring(xml_bytes)
        namespaces = {
            "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
        }

        blocks: list[str] = []
        for selector in (".//text:h", ".//text:p"):
            for node in root.findall(selector, namespaces):
                text = "".join(node.itertext()).strip()
                if text:
                    blocks.append(text)

        return "\n\n".join(blocks).strip()

    @staticmethod
    def _read_rtf(path: Path) -> str:
        if _which("pandoc"):
            try:
                result = subprocess.run(
                    ["pandoc", str(path), "-t", "plain"],
                    capture_output=True,
                    text=True,
                    check=True,
                )
                output = result.stdout.strip()
                if output:
                    return output
            except Exception:
                pass

        raw = DocumentLoader._read_plain_text(path)
        text = re.sub(r"\\par[d]?", "\n\n", raw)
        text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
        text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return html.unescape(text).strip()

    @staticmethod
    def _read_pages(path: Path) -> str:
        if _which("pandoc"):
            try:
                result = subprocess.run(
                    ["pandoc", str(path), "-t", "plain"],
                    capture_output=True,
                    text=True,
                    check=True,
                )
                output = result.stdout.strip()
                if output:
                    return output
            except Exception:
                pass

        raise UnsupportedFormatError(
            "Pages files are not reliable on Linux. Export them from Pages as DOCX first."
        )

    @staticmethod
    def _read_html(path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise UnsupportedFormatError("HTML support requires beautifulsoup4.") from exc

        raw = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")
        text = soup.get_text("\n\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise UnsupportedFormatError("PDF support is unavailable in this build.") from exc

        reader = PdfReader(str(path))
        pages: list[str] = []

        for page in reader.pages:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append(text)

        return "\n\n".join(pages).strip()


def _which(binary: str) -> str | None:
    result = subprocess.run(
        ["sh", "-lc", f"command -v {binary}"],
        capture_output=True,
        text=True,
    )
    path = result.stdout.strip()
    return path or None
